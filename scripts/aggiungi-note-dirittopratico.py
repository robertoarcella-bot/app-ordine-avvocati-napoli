"""
Aggiorna il Manuale (docx + md) aggiungendo la sezione "Note di Udienza —
Diritto Pratico" SENZA rigenerare l'intero documento, preservando così le
modifiche manuali dell'utente (in particolare il ridimensionamento delle
immagini per evitare che didascalie e testi vadano alla pagina successiva).

Strategia:
- Apre il docx esistente con python-docx
- Cerca il paragrafo footer "— Fine del Manuale —" e inserisce le nuove
  sezioni PRIMA (con un page break)
- Aggiorna anche il file .md aggiungendo la sezione in coda prima del footer

Idempotente: se la sezione "Note di Udienza" è già presente, non duplica.
"""
from pathlib import Path
import re

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

ROOT = Path(__file__).resolve().parents[1]
WORKSPACE = ROOT.parent
DOCX = WORKSPACE / "Manuale_App_OrdineAvvocatiNapoli.docx"
MD = WORKSPACE / "Manuale_App_OrdineAvvocatiNapoli.md"

PRIMARY_DARK = (0, 63, 127)
MUTED = (110, 117, 130)

NEW_SECTION_TITLE = "Note di Udienza — Diritto Pratico"


def _insert_paragraph_before(ref_para, text="", style=None):
    """Inserisce un nuovo paragrafo PRIMA di ref_para. Restituisce il nuovo paragrafo."""
    new_p = OxmlElement("w:p")
    ref_para._p.addprevious(new_p)
    from docx.text.paragraph import Paragraph
    para = Paragraph(new_p, ref_para._parent)
    if style:
        para.style = style
    if text:
        para.add_run(text)
    return para


def _insert_page_break_before(ref_para):
    """Inserisce un page break PRIMA del paragrafo di riferimento."""
    p = _insert_paragraph_before(ref_para)
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def add_heading_before(ref, text, level=1):
    p = _insert_paragraph_before(ref, style=f"Heading {level}")
    for run in p.runs:
        run.font.color.rgb = RGBColor(*PRIMARY_DARK)
    if not p.runs:
        r = p.add_run(text)
        r.font.color.rgb = RGBColor(*PRIMARY_DARK)
    else:
        # Lo style ha messo il run vuoto; cancello e aggiungo il testo
        for r in p.runs:
            r.text = text
            r.font.color.rgb = RGBColor(*PRIMARY_DARK)
    return p


def add_para_before(ref, text, italic=False, bold=False, color=None, size=None):
    p = _insert_paragraph_before(ref)
    # Supporta markdown ** per il bold
    parts = text.split("**")
    bold_state = False
    for part in parts:
        if part:
            r = p.add_run(part)
            r.bold = bold or bold_state
            if italic:
                r.italic = True
            if color:
                r.font.color.rgb = RGBColor(*color)
            if size:
                r.font.size = Pt(size)
        bold_state = not bold_state
    return p


def add_bullet_before(ref, text):
    p = _insert_paragraph_before(ref, style="List Bullet")
    parts = text.split("**")
    bold = False
    for part in parts:
        if part:
            r = p.add_run(part)
            r.bold = bold
        bold = not bold
    return p


# ===================================================================
# UPDATE DOCX
# ===================================================================
def update_docx():
    if not DOCX.exists():
        print(f"  [SKIP] {DOCX} non esiste")
        return False
    doc = Document(str(DOCX))

    # Cerca se la sezione esiste già (idempotenza)
    for para in doc.paragraphs:
        if NEW_SECTION_TITLE.lower() in para.text.lower():
            print(f"  [SKIP DOCX] Sezione '{NEW_SECTION_TITLE}' già presente")
            return False

    # Trova il footer "— Fine del Manuale —"
    target = None
    for para in doc.paragraphs:
        if "fine del manuale" in para.text.lower():
            target = para
            break
    if target is None:
        # Fallback: usa l'ultimo paragrafo
        target = doc.paragraphs[-1]
        print("  [INFO DOCX] Footer non trovato, inserisco in fondo")

    # Inserisci sezione PRIMA del footer
    _insert_page_break_before(target)

    add_heading_before(target, "21. Note di Udienza — Diritto Pratico", level=1)

    add_para_before(target,
        "L'app integra il servizio Note di Udienza Diritto Pratico "
        "(note.dirittopratico.it), uno strumento web ideato per la "
        "verbalizzazione delle note di udienza nell'era del Processo Civile "
        "Telematico. Permette di compilare comode note di udienza pronte da "
        "depositare o consultare in udienza dal proprio smartphone."
    )

    add_para_before(target,
        "L'integrazione avviene tramite il browser interno dell'app: il sito "
        "è già nativamente responsive e ottimizzato per dispositivi mobili, "
        "non contiene pubblicità, e la sessione utente resta attiva nel browser "
        "interno tra aperture successive."
    )

    add_heading_before(target, "Come accedere", level=2)
    add_bullet_before(target,
        "**Dal menu hamburger** → voce **Note di Udienza (Diritto Pratico)** — "
        "apre direttamente lo strumento nel browser interno."
    )
    add_bullet_before(target,
        "**Dalla sezione Strumenti** → categoria **Servizi di interesse "
        "generale** → voce **Note di Udienza — Diritto Pratico** — "
        "stesso comportamento."
    )

    add_heading_before(target, "Caratteristiche principali", level=2)
    for x in [
        "Compilazione note di udienza secondo gli schemi del PCT.",
        "Layout adattivo automatico in base alla larghezza del telefono.",
        "Nessuna pubblicità: l'utente lavora senza distrazioni.",
        "Sessione persistente: il browser interno dell'app conserva il lavoro "
        "non salvato tra aperture successive.",
        "Possibilità di esportare/condividere il risultato (a seconda delle "
        "funzioni offerte dal sito).",
    ]:
        add_bullet_before(target, x)

    add_para_before(target,
        "Lo strumento è offerto da Diritto Pratico ed è esterno all'app: "
        "eventuali aggiornamenti delle funzionalità, layout o limiti d'uso "
        "dipendono direttamente dal gestore del sito.",
        italic=True, color=MUTED
    )

    doc.save(str(DOCX))
    print(f"  [OK DOCX] Aggiornato {DOCX} ({DOCX.stat().st_size / 1024:.1f} KB)")
    return True


# ===================================================================
# UPDATE MD
# ===================================================================
def update_md():
    if not MD.exists():
        print(f"  [SKIP] {MD} non esiste")
        return False
    txt = MD.read_text(encoding="utf-8")
    if NEW_SECTION_TITLE.lower() in txt.lower():
        print(f"  [SKIP MD] Sezione '{NEW_SECTION_TITLE}' già presente")
        return False

    new_block = """
---

## 21. Note di Udienza — Diritto Pratico

L'app integra il servizio Note di Udienza Diritto Pratico (note.dirittopratico.it), uno strumento web ideato per la verbalizzazione delle note di udienza nell'era del Processo Civile Telematico. Permette di compilare comode note di udienza pronte da depositare o consultare in udienza dal proprio smartphone.

L'integrazione avviene tramite il browser interno dell'app: il sito è già nativamente responsive e ottimizzato per dispositivi mobili, non contiene pubblicità, e la sessione utente resta attiva nel browser interno tra aperture successive.

### Come accedere

- **Dal menu hamburger** → voce **Note di Udienza (Diritto Pratico)** — apre direttamente lo strumento nel browser interno.
- **Dalla sezione Strumenti** → categoria **Servizi di interesse generale** → voce **Note di Udienza — Diritto Pratico** — stesso comportamento.

### Caratteristiche principali

- Compilazione note di udienza secondo gli schemi del PCT.
- Layout adattivo automatico in base alla larghezza del telefono.
- Nessuna pubblicità: l'utente lavora senza distrazioni.
- Sessione persistente: il browser interno dell'app conserva il lavoro non salvato tra aperture successive.
- Possibilità di esportare/condividere il risultato (a seconda delle funzioni offerte dal sito).

> ℹ️  Lo strumento è offerto da Diritto Pratico ed è esterno all'app: eventuali aggiornamenti delle funzionalità, layout o limiti d'uso dipendono direttamente dal gestore del sito.
"""

    # Cerca il footer "— Fine del Manuale —" e inserisci PRIMA
    footer_pattern = re.compile(r"(\n+\*[^*]*[Ff]ine del [Mm]anuale[^*]*\*)", re.MULTILINE)
    m = footer_pattern.search(txt)
    if m:
        new_txt = txt[:m.start()] + new_block + txt[m.start():]
    else:
        # Fallback: append in coda
        new_txt = txt.rstrip() + "\n" + new_block + "\n"

    MD.write_text(new_txt, encoding="utf-8")
    print(f"  [OK MD]   Aggiornato {MD} ({MD.stat().st_size / 1024:.1f} KB)")
    return True


if __name__ == "__main__":
    import sys
    print("Aggiornamento manuale (.docx + .md) con sezione Note Diritto Pratico...")
    only_md = "--only-md" in sys.argv
    only_docx = "--only-docx" in sys.argv
    if not only_md:
        try:
            update_docx()
        except PermissionError:
            print("  [ERR DOCX] File aperto in Word — chiudilo e rilancia lo script")
    if not only_docx:
        update_md()
    print("Fatto.")
