"""
Genera il Manuale Utente dell'app Ordine Avvocati Napoli (v1.1.0).

Output (entrambi):
  - ../Manuale_App_OrdineAvvocatiNapoli.docx  (Microsoft Word)
  - ../Manuale_App_OrdineAvvocatiNapoli.md    (Markdown, leggibile su GitHub)

Le immagini sono reali screenshot dell'app (cartella FOTO_Manuale del workspace).

Vincoli di impaginazione applicati nel DOCX:
  - Larghezza immagine 7 cm (~50% pagina A4) per non occupare troppo spazio.
  - keep_with_next sul paragrafo immediatamente precedente all'immagine, così
    l'immagine non si separa dal testo che la introduce.
  - keep_with_next sul paragrafo dell'immagine e su quello della didascalia,
    così non si producono orfane fra immagine, caption e descrizione.
"""
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============== Setup percorsi ==============
ROOT = Path(__file__).resolve().parents[1]
WORKSPACE = ROOT.parent
PHOTO_SRC = WORKSPACE / "FOTO_Manuale"
IMAGES_DIR = ROOT / "scripts" / "manual-images"
IMAGES_DIR.mkdir(parents=True, exist_ok=True)

OUT_DOCX = WORKSPACE / "Manuale_App_OrdineAvvocatiNapoli.docx"
OUT_MD = WORKSPACE / "Manuale_App_OrdineAvvocatiNapoli.md"

# Mapping screenshot v1.1.0 → nome semantico
PHOTO_MAPPING = [
    ("WhatsApp Image 2026-05-03 at 17.01.55 (6).jpeg", "home.jpeg"),                       # Home con tile
    ("WhatsApp Image 2026-05-03 at 17.01.55 (5).jpeg", "news.jpeg"),                       # News con cards
    ("WhatsApp Image 2026-05-03 at 17.01.55 (4).jpeg", "sito.jpeg"),                       # Sito dell'Ordine
    ("WhatsApp Image 2026-05-03 at 17.01.55 (3).jpeg", "documenti.jpeg"),                  # Documenti
    ("WhatsApp Image 2026-05-03 at 17.01.55 (2).jpeg", "strumenti_generali.jpeg"),         # Strumenti — Servizi generali (con Codici Italiani)
    ("WhatsApp Image 2026-05-03 at 17.01.55.jpeg",     "strumenti_civile.jpeg"),           # Strumenti — Civile
    ("WhatsApp Image 2026-05-03 at 17.01.55 (1).jpeg", "codici.jpeg"),                     # Codici Italiani — home
    ("WhatsApp Image 2026-05-03 at 17.01.54 (3).jpeg", "tribunale.jpeg"),                  # Tribunale Napoli (magistrati con Torre/Piano)
    ("WhatsApp Image 2026-05-03 at 17.01.54 (2).jpeg", "processo_telematico.jpeg"),        # Processo Telematico
    ("WhatsApp Image 2026-05-03 at 17.01.54 (1).jpeg", "consiglio.jpeg"),                  # Consiglio dell'Ordine
    ("WhatsApp Image 2026-05-03 at 17.01.54.jpeg",     "commissione.jpeg"),                # Commissione Informatica
]


def copy_photos():
    """Copia gli screenshot nella cartella manual-images con nomi semantici."""
    paths = {}
    for src_name, dst_name in PHOTO_MAPPING:
        src = PHOTO_SRC / src_name
        if not src.exists():
            print(f"  [SKIP] {src_name} non trovato")
            continue
        dst = IMAGES_DIR / dst_name
        shutil.copy2(src, dst)
        paths[dst_name.replace(".jpeg", "")] = dst
        print(f"  [OK] {dst_name}")
    return paths


# ============== Palette colori ==============
PRIMARY_DARK = (0, 63, 127)
PRIMARY = (0, 102, 204)
MUTED = (110, 117, 130)


# ============== Modello del manuale ==============
# Lista di elementi: ogni elemento è (tipo, payload).
# Tipi: H1, H2, H3, P, BULLET, NUM, IMG (path, caption opzionale),
#       NOTE (paragrafo evidenziato), PAGEBREAK, QUOTE
def build_content(images):
    C = []  # lista di elementi
    h1 = lambda t: C.append(("H1", t))
    h2 = lambda t: C.append(("H2", t))
    h3 = lambda t: C.append(("H3", t))
    p = lambda t: C.append(("P", t))
    b = lambda t: C.append(("BULLET", t))
    n = lambda t: C.append(("NUM", t))
    img = lambda key, cap=None: C.append(("IMG", (images.get(key), cap)))
    note = lambda t: C.append(("NOTE", t))
    pb = lambda: C.append(("PAGEBREAK", None))
    quote = lambda t: C.append(("QUOTE", t))

    # ===================================================================
    # FRONTESPIZIO
    # ===================================================================
    C.append(("TITLE", "Manuale d'uso"))
    C.append(("SUBTITLE", "App Ordine Avvocati Napoli"))
    C.append(("SUB2", "Consiglio dell'Ordine degli Avvocati di Napoli"))
    C.append(("SUB3", "Versione 1.1.1 — maggio 2026"))
    img("home", "Schermata principale dell'app")
    C.append(("CREDIT", "Autore: Avv. Roberto Arcella"))
    C.append(("CREDIT_SUB", "Idea e collaborazione: Commissione Informatica del COA Napoli"))
    pb()

    # ===================================================================
    # SOMMARIO
    # ===================================================================
    h1("Sommario")
    sezioni = [
        "1. Introduzione",
        "2. Novità della versione 1.1.0",
        "3. Installazione",
        "4. Schermata principale (Home)",
        "5. Menu di navigazione",
        "6. Consiglio dell'Ordine",
        "7. News dal Consiglio",
        "8. News dagli Uffici Giudiziari",
        "9. Sito Ordine Avvocati",
        "10. Albo Avvocati Napoli e Albo Nazionale",
        "11. Documenti",
        "12. Strumenti",
        "13. Codici Italiani",
        "14. Dislocazione Aule e Uffici",
        "    14.1 Tribunale di Napoli",
        "    14.2 Corte d'Appello di Napoli",
        "    14.3 Cerca per autorità e sezione",
        "    14.4 Recapiti uffici giudiziari del distretto",
        "    14.5 Sezione Lavoro, Aule Penali, Calendario GdP",
        "15. Processo Telematico",
        "    15.1 Notifiche di nuovi avvisi PST",
        "    15.2 Fonti di news ufficiali",
        "    15.3 Strumenti operativi PCT",
        "16. Riconosco",
        "17. Area Riservata Consiglieri",
        "18. Commissione Informatica",
        "19. Info & Crediti",
        "20. Funzionalità trasversali",
        "21. Privacy e sicurezza",
        "22. Risoluzione problemi",
    ]
    for s in sezioni:
        n(s)
    pb()

    # ===================================================================
    # 1. INTRODUZIONE
    # ===================================================================
    h1("1. Introduzione")
    p("L'app Ordine Avvocati Napoli è uno strumento informativo e operativo "
      "realizzato per gli iscritti al Consiglio dell'Ordine degli Avvocati di Napoli. "
      "Concentra in un unico punto le funzioni più frequentemente utili al "
      "professionista forense: news del Consiglio, news degli Uffici Giudiziari "
      "di Napoli, accesso rapido al sito istituzionale, modulistica, calcolatori "
      "e strumenti operativi, codici italiani consultabili offline, informazioni "
      "sui processi telematici (PCT, PPT, amministrativo e tributario), accesso "
      "all'area riservata, dislocazione magistrati, aule d'udienza e cancellerie "
      "del Tribunale e della Corte d'Appello.")
    p("L'app è stata sviluppata dall'Avv. Roberto Arcella su idea e con la "
      "collaborazione della Commissione Informatica del Consiglio dell'Ordine "
      "degli Avvocati di Napoli.")
    h2("Caratteristiche principali")
    for x in [
        "Funziona prevalentemente offline: i calcolatori, gli strumenti, "
        "i codici italiani, l'elenco di magistrati e cancellerie, i contatti "
        "del Consiglio, la composizione del Consiglio e della Commissione "
        "Informatica sono incorporati nel pacchetto e non richiedono connessione.",
        "News in tempo reale dal sito ufficiale del Consiglio (via API REST "
        "WordPress), con cache locale (le ultime news sono consultabili "
        "anche offline).",
        "News aggregate da Tribunale di Napoli e Corte d'Appello di Napoli, "
        "presentate in ordine cronologico.",
        "Aggregatore di notizie sui processi telematici da fonti istituzionali: "
        "Ministero della Giustizia (PST), Giustizia Amministrativa, "
        "Dipartimento della Giustizia Tributaria.",
        "Notifiche automatiche di nuovi avvisi PST (background fetch periodico, "
        "previa autorizzazione dell'utente).",
        "Codici italiani offline (Costituzione, codice civile, codice penale, "
        "codici di procedura, CAD, Codice del processo amministrativo) "
        "ricercabili anche cross-codice e con copia/condivisione articoli su WhatsApp.",
        "Accesso integrato all'area riservata Consiglieri tramite il browser "
        "di sistema, con flusso a 2 step esplicito.",
        "Tema istituzionale conforme alle linee guida AgID per i siti della "
        "Pubblica Amministrazione italiana.",
        "Supporto del tema scuro automatico in base alle preferenze del telefono.",
    ]:
        b(x)
    pb()

    # ===================================================================
    # 2. NOVITÀ v1.1.0
    # ===================================================================
    h1("2. Novità della versione 1.1.0")
    p("Questa edizione del manuale documenta le novità introdotte con la "
      "versione 1.1.0 dell'app, rilasciata a maggio 2026. Le principali "
      "aree interessate sono tre.")

    h2("Codici Italiani offline")
    p("È stata integrata, completamente in app, una raccolta dei principali "
      "codici italiani: Costituzione, Codice civile, Disposizioni di attuazione "
      "del codice civile, Codice di procedura civile, Disposizioni di attuazione "
      "del codice di procedura civile, Codice penale, Codice di procedura penale, "
      "Norme di attuazione del codice di procedura penale, Regolamento di "
      "esecuzione del codice di procedura penale, Codice dell'amministrazione "
      "digitale (CAD) e Codice del processo amministrativo (CPA). Tutti i "
      "testi sono navigabili e ricercabili senza connessione e ogni articolo "
      "può essere copiato o condiviso direttamente su WhatsApp.")

    h2("Aule giudici Tribunale di Napoli (Torre A)")
    p("L'elenco dei magistrati del Tribunale di Napoli ora indica per ciascun "
      "giudice la Torre e il Piano della relativa aula d'udienza. Toccando un "
      "magistrato si apre un riquadro a tutto schermo con tutte le informazioni "
      "in formato leggibile, compresi (per i giudici della Sezione Lavoro) i "
      "giorni d'udienza tipici.")

    h2("Nuove pagine di ricerca uffici")
    p("È stata aggiunta una pagina di ricerca per autorità (Tribunale o Corte "
      "d'Appello) e sezione, in cui si possono visualizzare in un solo colpo "
      "d'occhio tutti i magistrati assegnati a una determinata sezione con il "
      "loro piano. È stata inoltre aggiunta una pagina con i recapiti completi "
      "(indirizzo, telefono, e-mail, fax) di tutti gli uffici del distretto "
      "della Corte d'Appello di Napoli — circa 80 voci, ricercabili.")
    pb()

    # ===================================================================
    # 3. INSTALLAZIONE
    # ===================================================================
    h1("3. Installazione")
    p("Nella prima fase l'app viene distribuita come file APK Android, da "
      "installare manualmente. Il file installer si chiama "
      "OrdineAvvocatiNapoli-1.1.1-debug.apk.")
    h2("Procedura su Android")
    for x in [
        "Trasferire l'APK sul telefono (e-mail, Google Drive, USB, link diretto).",
        "Aprire il file. Android chiederà di abilitare l'installazione da fonti "
        "sconosciute per il browser o file manager: confermare.",
        "Avviare l'installazione. L'app appare come Ordine Avvocati Napoli con "
        "icona blu istituzionale.",
        "Al primo avvio compare lo splash screen istituzionale (logo COA su "
        "fondo crema), poi la schermata principale.",
    ]:
        b(x)
    p("All'avvio non sono richieste registrazioni: l'app è subito utilizzabile "
      "in tutta la parte ad accesso libero. L'area riservata Consiglieri "
      "richiede le credenziali del sito istituzionale del Consiglio.")
    note("Versione iOS (iPhone): per ora la build iOS è gestita su un fork "
         "separato della repo (vedi documento BUILD_iOS.md sul repository GitHub). "
         "Richiede un Mac con Xcode e — per la distribuzione su TestFlight o "
         "App Store — un Apple Developer Program (~99 €/anno).")
    pb()

    # ===================================================================
    # 4. HOME
    # ===================================================================
    h1("4. Schermata principale (Home)")
    p("La home presenta in alto il logo del Consiglio e la denominazione "
      "completa, e a seguire una griglia di riquadri rapidi che danno accesso "
      "alle aree principali dell'app.")
    img("home", "Home — logo del Consiglio e tile di accesso rapido")
    h2("Riquadri disponibili")
    tiles = [
        ("News", "Ultime news del Consiglio dell'Ordine"),
        ("Sito", "Sezioni del sito istituzionale (in WebView interna)"),
        ("Documenti", "Modulistica, Albo, Trasparenza, Verbali"),
        ("Strumenti", "Oltre 20 calcolatori e strumenti professionali, codici italiani"),
        ("Dislocazione Aule e Uffici", "Tribunale, Corte d'Appello, ricerca per sezione, recapiti distretto"),
        ("Processo Telematico", "Notizie da PST, GA, Giustizia Tributaria"),
        ("Area Riservata", "Accesso area riservata Consiglieri (browser di sistema)"),
        ("Riconosco", "Sistema di identità digitale dei consigli forensi"),
    ]
    for name, desc in tiles:
        b(f"**{name}** — {desc}")
    p("In ogni schermata è disponibile l'icona ☰ in alto a sinistra che apre "
      "il menu laterale con tutte le voci di navigazione (vedi sezione successiva).")
    pb()

    # ===================================================================
    # 5. MENU
    # ===================================================================
    h1("5. Menu di navigazione")
    p("Toccando l'icona ☰ in alto a sinistra (o scorrendo dal bordo sinistro "
      "dello schermo) si apre il menu laterale con tutte le voci di "
      "navigazione dell'app, in ordine logico di consultazione.")
    h2("Voci del menu")
    for v in [
        "**Home** — Torna alla schermata principale.",
        "**Consiglio dell'Ordine** — Composizione del Consiglio in carica (Presidenza + Consiglieri).",
        "**News dal Consiglio** — Tutte le news del Consiglio, con ricerca e infinite scroll.",
        "**News dagli Uffici Giudiziari** — News aggregate da Tribunale e Corte d'Appello di Napoli.",
        "**Sito Ordine Avvocati** — Indice delle sezioni del sito istituzionale.",
        "**Albo Avvocati Napoli** — Apre l'albo iscritti al COA Napoli direttamente nel browser interno.",
        "**Albo Nazionale Avvocati** — Apre il portale ricerca avvocati del CNF in browser di sistema.",
        "**Documenti** — Modulistica e documenti pubblici di rilievo.",
        "**Strumenti** — Calcolatori, strumenti operativi e codici italiani.",
        "**Dislocazione Aule e Uffici** — Tribunale, Corte d'Appello, ricerca per sezione, recapiti distretto, calendari e canali.",
        "**Processo Telematico** — Notizie e strumenti per i processi telematici, notifiche di nuovi avvisi.",
        "**Area Riservata** — Accesso area riservata Consiglieri (browser di sistema).",
        "**Commissione Informatica** — Componenti della Commissione che ha ideato l'app.",
        "**Info & Crediti** — Contatti del Consiglio (offline) e crediti dell'app.",
    ]:
        b(v)
    pb()

    # ===================================================================
    # 6. CONSIGLIO
    # ===================================================================
    h1("6. Consiglio dell'Ordine")
    p("Mostra la composizione completa del Consiglio dell'Ordine in carica, "
      "suddivisa in due sezioni: l'Ufficio di Presidenza (Presidente, "
      "Segretario, Tesoriera, Vice Presidenti) e i Consiglieri, in ordine "
      "di pubblicazione sul sito istituzionale.")
    img("consiglio", "Composizione del Consiglio dell'Ordine — Ufficio di Presidenza e Consiglieri")
    p("Per ciascun componente è indicato il nome con il prefisso Avv. e, "
      "sotto, la carica ricoperta. Una piccola icona evidenzia "
      "l'appartenenza all'Ufficio di Presidenza.")
    note("Questa pagina contiene dati offline-first incorporati nell'app. "
         "Vanno aggiornati ad ogni rinnovo del Consiglio o sostituzione di "
         "un componente, ricompilando l'app.")
    pb()

    # ===================================================================
    # 7. NEWS COA
    # ===================================================================
    h1("7. News dal Consiglio")
    p("Aggregatore delle news pubblicate sul sito istituzionale del Consiglio. "
      "Le news sono recuperate tramite l'API REST ufficiale di WordPress "
      "del sito (no scraping), con cache locale di 30 minuti per ridurre "
      "il consumo di dati.")
    img("news", "Lista news con barra di ricerca e cards riassuntive")
    h2("Funzioni")
    for x in [
        "Ricerca testuale all'interno delle news (campo cerca in alto).",
        "Pull-to-refresh per forzare il ricaricamento.",
        "Scroll infinito: scorrendo in basso vengono caricate le news più datate.",
        "Tocco su una news per leggere il contenuto completo.",
        "Pulsante condividi per inviare la news via WhatsApp, e-mail o altre app.",
        "Pulsante \"Apri sul sito\" per visualizzare la news direttamente sul "
        "sito istituzionale.",
    ]:
        b(x)
    pb()

    # ===================================================================
    # 8. NEWS UFFICI
    # ===================================================================
    h1("8. News dagli Uffici Giudiziari")
    p("Sezione dedicata che aggrega in un unico flusso cronologico le news "
      "pubblicate dai siti istituzionali del Tribunale di Napoli e della "
      "Corte d'Appello di Napoli. Ogni news riporta un chip con l'indicazione "
      "dell'ufficio di provenienza, la data, il titolo e l'estratto.")
    h2("Funzioni")
    for x in [
        "Aggregazione in tempo reale dei due flussi, ordinati per data discendente.",
        "Cache locale di 30 minuti.",
        "Pull-to-refresh per forzare il ricaricamento.",
        "Tocco su una news per aprirla nel browser di sistema sul sito di origine.",
        "Robusto a fallimenti parziali: se una delle due fonti è temporaneamente "
        "indisponibile, l'altra continua a essere mostrata.",
    ]:
        b(x)
    pb()

    # ===================================================================
    # 9. SITO
    # ===================================================================
    h1("9. Sito Ordine Avvocati")
    p("Indice delle sezioni del sito istituzionale rese accessibili "
      "direttamente dall'app, ciascuna in WebView interna ottimizzata per "
      "dispositivi mobili.")
    img("sito", "Indice delle sezioni del sito istituzionale del COA")
    h2("Sezioni disponibili")
    for s in [
        "Componenti C.O.A.",
        "Commissioni",
        "Verbali sedute consiliari",
        "Modulistica",
        "Amministrazione Trasparente",
        "Protocolli d'intesa",
        "Compiti e regolamenti",
        "Deposito negoziazioni assistite (sito esterno)",
        "Contatti",
    ]:
        b(s)
    pb()

    # ===================================================================
    # 10. ALBI
    # ===================================================================
    h1("10. Albo Avvocati Napoli e Albo Nazionale")
    h2("Albo Avvocati Napoli")
    p("Voce di menu autonoma che apre direttamente la sezione Albo ed Elenchi "
      "del sito ufficiale del COA Napoli (avvocati, praticanti, elenchi "
      "specialistici) all'interno del browser dell'app.")
    h2("Albo Nazionale Avvocati")
    p("Apre il portale di ricerca avvocati del Consiglio Nazionale Forense "
      "(consiglionazionaleforense.it/ricerca-avvocati). Per ragioni tecniche "
      "(il sito CNF blocca il caricamento in iframe tramite header X-Frame-Options) "
      "viene aperto nel browser di sistema, dove la sessione e l'esperienza "
      "utente sono pienamente garantite.")
    pb()

    # ===================================================================
    # 11. DOCUMENTI
    # ===================================================================
    h1("11. Documenti")
    p("Catalogo dei documenti pubblici di rilievo del Consiglio. Le voci con "
      "icona di download avviano il salvataggio del file nella cartella "
      "Documenti del telefono; quelle con icona di apertura esterna aprono la "
      "relativa pagina del sito.")
    img("documenti", "Catalogo documenti — modulistica, albo, regolamenti, verbali, trasparenza")
    h2("Voci disponibili")
    for s in [
        "Modulistica Ordine Professionale — moduli per iscrizioni, cancellazioni, certificati.",
        "Albo ed Elenchi — avvocati, praticanti, elenchi specialistici.",
        "Compiti e regolamenti — Statuto, regolamenti interni, codici deontologici.",
        "Verbali sedute consiliari — delibere e verbali del Consiglio.",
        "Amministrazione Trasparente — documenti ai sensi del d.lgs. 33/2013.",
    ]:
        b(s)
    pb()

    # ===================================================================
    # 12. STRUMENTI
    # ===================================================================
    h1("12. Strumenti")
    p("Catalogo di calcolatori, strumenti professionali e raccolte normative, "
      "raggruppati per area di competenza. Una piccola icona accanto al titolo "
      "indica se lo strumento funziona offline (☁) o richiede internet (📶).")
    img("strumenti_generali", "Strumenti — Servizi di interesse generale (con Codici Italiani)")

    h2("Servizi di interesse generale")
    for s in [
        "Parametri Forensi (D.M. 147/2022) — calcolatore parcelle.",
        "Calcolo Fattura Avvocato (CPA, IVA, ritenuta).",
        "Preventivo professionale.",
        "Contributo Unificato (Civile/Tributario/Amministrativo).",
        "Interessi Legali e Moratori (tasso legale, BCE).",
        "Procura alle liti — generatore.",
        "Note di Udienza — Diritto Pratico (verbalizzazione PCT, online).",
        "**Codici Italiani** — Costituzione, codici civile/penale e relative "
        "disposizioni di attuazione, CAD e Cod. Proc. Amministrativo, con "
        "copia/condivisione articoli su WhatsApp (vedi sezione successiva).",
    ]:
        b(s)

    img("strumenti_civile", "Strumenti — area Civile")
    h2("Civile")
    for s in [
        "Termini c.p.c. (Cartabia) — calcolatore termini processuali.",
        "Termini Esecuzioni Civili.",
        "Danno alla Persona — TUN 2025 (D.P.R. 12/2025).",
        "Danno Micropermanenti (metodo coefficienti).",
        "Contributo Unificato Famiglia.",
        "Piano Genitoriale — redazione assistita.",
        "Analizzatore atti D.M. 110/2023.",
        "Analisi Verbale Ricerca Beni.",
        "FAQ Patrocinio Spese Stato (Civile).",
    ]:
        b(s)

    h2("Penale")
    for s in [
        "Calcolo Prescrizione Penale (avanzato).",
        "Patrocinio S.S. Penale — Protocollo Napoli.",
        "PEC Uffici Giudiziari (PPT).",
    ]:
        b(s)

    note("Tutti gli strumenti tradizionalmente pubblicati su "
         "avvocatotelematico.studiolegalearcella.it sono inclusi nel pacchetto "
         "dell'app a beneficio degli iscritti.")
    pb()

    # ===================================================================
    # 13. CODICI ITALIANI (NUOVO v1.1)
    # ===================================================================
    h1("13. Codici Italiani")
    p("Toccando la voce \"Codici Italiani\" tra gli Strumenti si accede alla "
      "raccolta completa dei principali codici e testi normativi della "
      "legislazione italiana, completamente disponibili offline e ricercabili "
      "anche cross-codice tramite la barra di ricerca globale.")
    img("codici", "Codici Italiani — home con barra di ricerca globale")

    h2("Codici disponibili")
    p("L'elenco è organizzato secondo la consueta partizione delle aree del "
      "diritto, in modo da rendere la ricerca rapida e intuitiva:")
    for s in [
        "Costituzione della Repubblica Italiana.",
        "Codice civile.",
        "Disposizioni per l'attuazione del Codice civile e disposizioni transitorie.",
        "Codice di procedura civile.",
        "Disposizioni per l'attuazione del Codice di procedura civile e disposizioni transitorie.",
        "Codice penale.",
        "Codice di procedura penale.",
        "Norme di attuazione, di coordinamento e transitorie del Codice di procedura penale.",
        "Regolamento di esecuzione del Codice di procedura penale.",
        "Codice dell'amministrazione digitale (CAD) — D.Lgs. 7 marzo 2005, n. 82.",
        "Codice del processo amministrativo (CPA) — D.Lgs. 2 luglio 2010, n. 104, "
        "con i relativi Allegati 1 (Codice), 2 (Norme di attuazione), 3 (Norme "
        "transitorie) e 4 (Disposizioni abrogative e di coordinamento).",
    ]:
        b(s)

    h2("Funzioni di consultazione")
    for x in [
        "**Ricerca globale**: la barra di ricerca nella home cerca in tutti gli "
        "11 codici contemporaneamente per parola chiave, numero di articolo o titolo.",
        "**Sfoglia per codice**: ogni card apre il codice corrispondente con un "
        "indice degli articoli sul lato (sidebar) e il testo dell'articolo "
        "selezionato a tutta pagina.",
        "**Note di aggiornamento**: ogni articolo riporta in calce, in box "
        "dedicato, le note di aggiornamento Normattiva (es. \"AGGIORNAMENTO (Xa)\") "
        "con il riferimento alle leggi modificative.",
        "**Copia articolo**: il pulsante \"Copia\" inserisce il testo "
        "dell'articolo (numero, rubrica, contenuto) negli appunti del telefono, "
        "pronto per essere incollato in qualsiasi documento.",
        "**Condivisione su WhatsApp**: il pulsante WhatsApp apre direttamente "
        "una chat con il testo dell'articolo precompilato.",
    ]:
        b(x)
    note("CAD e CPA sono stati generati automaticamente a partire dagli XML "
         "AkomaNtoso pubblicati da Normattiva.it, lo stesso formato standard "
         "utilizzato dagli altri codici. Tutti i testi risultano aggiornati alle "
         "ultime versioni vigenti pubblicate.")
    pb()

    # ===================================================================
    # 14. AULE UDIENZE
    # ===================================================================
    h1("14. Dislocazione Aule e Uffici")
    p("Sezione dedicata alla logistica delle udienze, alla dislocazione degli "
      "uffici giudiziari di Napoli e ai recapiti degli uffici dell'intero "
      "distretto della Corte d'Appello. La home della sezione presenta in alto "
      "i due grandi blocchi (Tribunale e Corte d'Appello), poi la voce di "
      "ricerca per autorità e sezione, poi i recapiti del distretto, infine "
      "i calendari e i canali specifici.")

    h2("14.1 Tribunale di Napoli")
    p("Pagina dedicata con quattro tab interne (segments) e ricerca; "
      "l'intestazione è ora in bianco su blu istituzionale per assicurare "
      "la massima leggibilità.")
    img("tribunale", "Tribunale di Napoli — magistrati con Torre e Piano dell'aula d'udienza")
    for x in [
        "**Magistrati** — circa 350 magistrati raggruppati per sezione, con "
        "filtro per area giurisdizione (Civile / Lavoro / Penale) e ricerca "
        "testuale per cognome, nome o sezione. Per ogni magistrato sono "
        "indicati nome, cognome, ruolo (Presidente di Sezione, Giudice, "
        "Giudice Onorario, GIP/GUP, ecc.) e — novità della v1.1 — la "
        "**Torre e il Piano** dell'aula d'udienza (Torre A, piani 7-22 a "
        "seconda della sezione). Toccando un magistrato si apre una scheda "
        "a tutto schermo con tutte le informazioni in formato leggibile, "
        "compresi (per i giudici della Sezione Lavoro) i giorni d'udienza tipici.",
        "**Uffici** — gli uffici amministrativi (Presidenza, Dirigenza, "
        "Personale, Recupero Crediti, Spese pagate dall'Erario, Economato, "
        "Archivio, Corpi di Reato, Centralino) con dislocazione su "
        "torre/piano, responsabile, telefono cliccabile, email cliccabile, PEC.",
        "**Dislocazione** — mappa logica dei settori del Tribunale per "
        "torre (A/B/C) e range di piani (es. Settore Civile su torre A "
        "piani 6-22, Settore Penale GIP/GUP su torre B piani 12-16, "
        "Sezioni Riesame torre C piano 5).",
        "**Info** — dati identificativi (sede, telefono, PEC, email).",
    ]:
        b(x)

    h2("14.2 Corte d'Appello di Napoli")
    p("Pagina dedicata con cinque tab interne e ricerca:")
    for x in [
        "**Magistrati** — i magistrati delle sezioni civili, lavoro "
        "(5 unità), penali (6 sezioni) e Sezione Minorenni-Persona-Famiglia, "
        "con filtro per area e ricerca.",
        "**Uffici** — uffici e cancellerie con responsabile, telefono, "
        "email, PEC. Raggruppati per torre.",
        "**Vertici** — Presidente, Vicario, Coordinatori dei settori "
        "(Civile, Penale, Lavoro, Sezioni Assise Appello, Innovazione).",
        "**Info** — dati identificativi della Corte d'Appello.",
    ]:
        b(x)

    h2("14.3 Cerca per autorità e sezione")
    p("Pagina di ricerca, novità della v1.1, che permette di scegliere "
      "rapidamente l'autorità giudiziaria (Tribunale di Napoli oppure Corte "
      "d'Appello di Napoli) e poi la sezione di interesse, ottenendo l'elenco "
      "completo dei magistrati assegnati a quella sezione, con il loro piano "
      "ed eventuali giorni d'udienza. Lo strumento è particolarmente utile "
      "quando si conosce la sezione di un fascicolo ma non il nome del "
      "magistrato, oppure quando si vuole avere un colpo d'occhio sulla "
      "composizione di un'intera sezione.")

    h2("14.4 Recapiti uffici giudiziari del distretto")
    p("Pagina, anch'essa novità della v1.1, con i recapiti completi (indirizzo, "
      "CAP, telefono, e-mail istituzionale, fax) di tutti gli uffici giudiziari "
      "del distretto della Corte d'Appello di Napoli — circa 80 voci, "
      "ricercabili per nome, comune, circondario o tipologia.")
    h3("Uffici inclusi")
    for x in [
        "Uffici distrettuali: Corte d'Appello, Procura Generale, Tribunale "
        "per i Minorenni, Procura presso il Tribunale per i Minorenni, "
        "Tribunale di Sorveglianza, Uffici di Sorveglianza (Avellino, Napoli, "
        "Santa Maria Capua Vetere), Tribunale Regionale delle Acque Pubbliche, "
        "Corte di Assise d'Appello.",
        "Circondario di Ariano Irpino: Tribunale, Procura, Giudici di Pace.",
        "Circondario di Avellino: Tribunale, Sezione Distaccata, Procura, "
        "Giudici di Pace, Corte di Assise.",
        "Circondario di Benevento: Tribunale, Sezioni Distaccate, Procura, "
        "Giudici di Pace, Corte di Assise.",
        "Circondario di Giugliano in Campania: Tribunale, Sezioni Distaccate "
        "(Pozzuoli, Marano), Procura, Giudici di Pace.",
        "Circondario di Napoli: Tribunale, Sezioni Distaccate (Portici, "
        "Afragola, Capri, Casoria, Frattamaggiore, Ischia), Procura, Giudici "
        "di Pace, Corte di Assise.",
        "Circondario di Nola: Tribunale, Procura, Giudici di Pace.",
        "Circondario di Sant'Angelo dei Lombardi: Tribunale, Procura, Giudici di Pace.",
        "Circondario di Santa Maria Capua Vetere: Tribunale, Sezioni Distaccate "
        "(Piedimonte Matese, Aversa, Carinola, Caserta, Marcianise), Procura, "
        "Giudici di Pace, Corte di Assise.",
        "Circondario di Torre Annunziata: Tribunale, Sezioni Distaccate "
        "(Sorrento, Castellammare di Stabia, Gragnano, Torre del Greco), "
        "Procura, Giudici di Pace.",
    ]:
        b(x)
    p("Toccando un indirizzo si apre Google Maps con la posizione; toccando "
      "un numero di telefono parte la chiamata; toccando un'e-mail si apre il "
      "client di posta del telefono. I dati sono offline-first: una volta "
      "installata l'app, sono disponibili anche senza connessione.")

    h2("14.5 Sezione Lavoro, Aule Penali, Calendario GdP")
    for x in [
        "**Aule Sezione Lavoro Napoli** — calendario delle udienze dei "
        "giudici della Sezione Lavoro del Tribunale, in formato a card "
        "responsive con sezione, piano e giorni di udienza in evidenza. "
        "Filtri per giorno della settimana (auto-seleziona il giorno corrente) "
        "e per sezione (I, II, III). Ricerca per cognome con autocompletamento.",
        "**Aule Penali Napoli** — apertura del canale Telegram "
        "@AulePenaliNapoli con segnalazioni in tempo reale. Richiede l'app "
        "Telegram installata.",
        "**Calendario GdP Napoli 2026** — calendario delle udienze del "
        "Giudice di Pace di Napoli per l'anno corrente, completamente offline.",
    ]:
        b(x)
    pb()

    # ===================================================================
    # 15. PROCESSO TELEMATICO
    # ===================================================================
    h1("15. Processo Telematico")
    p("Aggregatore di notizie istituzionali e strumenti operativi sui processi "
      "telematici (civile, penale, amministrativo, tributario), con un sistema "
      "di notifiche per i nuovi avvisi del PST.")
    img("processo_telematico", "Processo Telematico — toggle notifiche, fonti istituzionali e strumenti operativi")

    h2("15.1 Notifiche di nuovi avvisi PST")
    p("In cima alla pagina è presente un toggle per attivare le notifiche "
      "automatiche di nuovi avvisi pubblicati sul Portale Servizi Telematici "
      "del Ministero della Giustizia. All'attivazione l'app richiede il "
      "permesso del sistema operativo per inviare notifiche.")
    p("Una volta attive, l'app verifica periodicamente in background "
      "(circa una volta all'ora; il sistema Android decide il timing preciso "
      "in base alle policy di risparmio batteria) e ad ogni apertura dell'app "
      "se ci sono nuovi avvisi rispetto a quelli già letti. In caso "
      "affermativo, schedula una notifica locale con il titolo della novità.")
    note("Su alcuni dispositivi con ottimizzazione batteria aggressiva "
         "(Xiaomi/MIUI, Huawei/EMUI) può essere necessario escludere l'app "
         "dalle impostazioni di risparmio energetico per ricevere le notifiche "
         "con regolarità in background.")

    h2("15.2 Fonti di news ufficiali")
    for x in [
        "**PST — Min. Giustizia**: avvisi su PCT, PPT, malfunzionamenti dei "
        "sistemi telematici (ordinari e penali). Scraping HTML.",
        "**Giustizia Amministrativa**: news del Consiglio di Stato e dei TAR "
        "(WebView interna per problemi SSL del sito).",
        "**Giustizia Tributaria (DGT MEF)**: feed RSS ufficiale con avvisi "
        "e rassegne sentenze.",
    ]:
        b(x)

    h2("15.3 Strumenti operativi PCT")
    for x in [
        "**Depositi CCII** — selettore atto-ruolo per i depositi nel Codice "
        "della Crisi d'Impresa.",
        "**Mappa XSD Depositi PCT** — schemi SICID/SIECIC + CCII.",
        "**SIGP — Consultazione Giudice di Pace** — sito legacy del Min. "
        "Giustizia, apre nel browser di sistema per maggiore usabilità.",
    ]:
        b(x)
    pb()

    # ===================================================================
    # 16. RICONOSCO
    # ===================================================================
    h1("16. Riconosco")
    p("Apre il portale Riconosco (riconosco.dcssrl.it), sistema di identità "
      "digitale degli iscritti agli ordini forensi italiani, in Chrome Custom "
      "Tabs di sistema. Questa scelta garantisce il corretto funzionamento "
      "del banner privacy del sito Riconosco, della sessione utente e di "
      "tutte le funzionalità avanzate.")

    # ===================================================================
    # 17. AREA RISERVATA
    # ===================================================================
    h1("17. Area Riservata Consiglieri")
    p("Accesso all'area riservata del sito istituzionale, riservata ai "
      "consiglieri e ai professionisti che vi sono abilitati.")
    p("La pagina propone un workflow esplicito a 2 step (in alto un avviso "
      "giallo che lo segnala), perché il sito ufficiale, per come è "
      "attualmente strutturato, non redirige automaticamente all'area "
      "Consiglieri dopo il login.")
    h2("Step 1 — Login")
    p("Tocca **Apri pagina di login** per aprire la pagina di login del sito "
      "istituzionale del Consiglio nel browser di sistema. Compila il modulo "
      "con le tue credenziali.")
    h2("Step 2 — Area Riservata Consiglieri")
    p("Una volta effettuato il login, torna in app e tocca il bottone "
      "evidenziato in blu **Vai all'Area Riservata Consiglieri**, che apre "
      "direttamente l'URL specifico della sezione (anch'esso visibile in "
      "chiaro nella card dell'app per maggiore trasparenza).")
    note("Le credenziali vengono inserite direttamente sul sito ufficiale "
         "del COA dentro il browser di sistema: l'app non le vede e non le "
         "memorizza. La sessione resta attiva nel browser di sistema anche "
         "tra aperture successive dell'app, e supporta nativamente i password "
         "manager del telefono.")
    pb()

    # ===================================================================
    # 18. COMMISSIONE
    # ===================================================================
    h1("18. Commissione Informatica")
    p("Pagina dedicata ai componenti della Commissione Informatica del "
      "Consiglio dell'Ordine degli Avvocati di Napoli, ideatrice e "
      "collaboratrice del progetto.")
    img("commissione", "Composizione della Commissione Informatica — Coordinatore e Componenti")
    p("L'elenco riporta:")
    for x in [
        "Il **Delegato all'informatica e all'innovazione** (Cons. Avv. "
        "Roberto Arcella).",
        "Il **Coordinatore** della Commissione (Avv. Leonardo Scinto).",
        "I **Componenti** in ordine alfabetico per cognome.",
    ]:
        b(x)
    pb()

    # ===================================================================
    # 19. INFO
    # ===================================================================
    h1("19. Info & Crediti")
    p("Pagina informativa con i dati di contatto del Consiglio dell'Ordine "
      "(disponibili anche offline) e i crediti dell'app.")
    h2("Contatti del Consiglio (offline)")
    for x in [
        "**Sede**: Centro Direzionale, Piazza Coperta — 80143 Napoli "
        "(tap → apre Google Maps)",
        "**Telefono**: +39 081 734 3737 (tap → avvia chiamata)",
        "**Fax**: +39 081 734 3010",
        "**Email**: segreteria@ordineavvocati.napoli.it (tap → client mail)",
        "**PEC**: segreteria@avvocatinapoli.legalmail.it",
        "**Orari segreteria**: Lunedì–Venerdì, 9.00–12.30",
        "**Codice Fiscale**: 80013690633",
        "**Codice Univoco fatturazione**: UF9L1M",
    ]:
        b(x)
    h2("Crediti")
    b("**Autore**: Avv. Roberto Arcella")
    b("**Idea e collaborazione**: Commissione Informatica COA Napoli")
    b("**Per conto di**: Consiglio dell'Ordine degli Avvocati di Napoli")
    pb()

    # ===================================================================
    # 20. FUNZIONALITÀ TRASVERSALI
    # ===================================================================
    h1("20. Funzionalità trasversali")
    h2("Tema scuro")
    p("L'app supporta automaticamente il tema scuro/chiaro in base alle "
      "preferenze di sistema del telefono. Anche le mini-webapp incluse si "
      "adattano al tema, fatta eccezione per i Codici Italiani che adottano "
      "il tema istituzionale del COA (blu e bianco) in modo costante per "
      "garantire la massima leggibilità del testo normativo.")
    h2("Funzionamento offline")
    p("La maggior parte degli strumenti funziona senza connessione: i "
      "calcolatori principali sono completamente autonomi (icona ☁), gli altri "
      "richiedono la prima apertura con internet (icona 📶). I contatti "
      "del Consiglio, la composizione del Consiglio e della Commissione "
      "Informatica, l'elenco dei magistrati e cancellerie del Tribunale e "
      "della Corte d'Appello, **i recapiti completi degli uffici del distretto** "
      "e **gli undici codici italiani** sono sempre disponibili offline. Le "
      "ultime news lette restano consultabili offline grazie alla cache.")
    h2("Apertura siti esterni")
    p("I siti istituzionali integrabili in iframe (PST, sito COA, sezioni) "
      "si aprono nel browser interno dell'app per un'esperienza coerente. "
      "I siti che richiedono comportamenti particolari (Riconosco, SIGP, "
      "CNF Albo Nazionale, area riservata) vengono aperti in Chrome Custom "
      "Tabs per garantire la corretta gestione di cookie, sessione e "
      "rendering.")
    h2("Pull-to-refresh")
    p("Nelle pagine di news (sia del Consiglio che degli Uffici Giudiziari "
      "che delle fonti del Processo Telematico), tirare verso il basso "
      "aggiorna la lista dei contenuti.")
    h2("Condivisione")
    p("Il pulsante condividi presente nei dettagli delle news permette di "
      "inviarle ad altre app del telefono (WhatsApp, e-mail, Telegram, "
      "ecc.). Anche le mini-webapp che producono testi (calcoli, modelli) "
      "possono inviare il risultato via WhatsApp grazie all'integrazione "
      "automatica. Per i Codici Italiani, ogni articolo dispone di un "
      "pulsante dedicato per copiarlo negli appunti e di uno per inviarlo "
      "direttamente su WhatsApp.")
    pb()

    # ===================================================================
    # 21. PRIVACY
    # ===================================================================
    h1("21. Privacy e sicurezza")
    p("L'app è progettata con attenzione alla privacy degli utenti:")
    for s in [
        "Le password dell'area riservata non vengono mai viste né memorizzate "
        "dall'app: il login avviene direttamente sul sito ufficiale dentro "
        "il browser di sistema.",
        "I dati personali eventualmente inseriti nei calcolatori restano sul "
        "telefono e non vengono trasmessi a server esterni.",
        "I cookie di sessione del sito ufficiale sono memorizzati nel browser "
        "interno dell'app, esattamente come in un browser normale.",
        "Non sono presenti tracker pubblicitari né strumenti di profilazione.",
        "Le news vengono recuperate dall'API REST ufficiale del sito del COA, "
        "che è pubblica e non richiede autenticazione.",
        "Le notifiche PST sono opzionali e attivabili solo previa esplicita "
        "autorizzazione dell'utente; l'app comunica al sistema operativo "
        "soltanto le credenziali tecniche per la schedulazione del task "
        "background.",
        "I codici italiani sono incorporati nell'app come testi normativi "
        "estratti da Normattiva.it: la consultazione non richiede connessione "
        "né tracciamento.",
    ]:
        b(s)
    pb()

    # ===================================================================
    # 22. RISOLUZIONE PROBLEMI
    # ===================================================================
    h1("22. Risoluzione problemi")
    h2("L'app non scarica le news")
    p("Verificare la connessione a internet. Se persiste, tirare la lista "
      "delle news verso il basso per forzare il refresh; in caso di errore "
      "di rete viene mostrato un messaggio con la possibilità di riprovare "
      "o di aprire il sito direttamente in browser.")
    h2("Una mini-app non scorre")
    p("Toccare brevemente lo schermo per attivare il gesture handler. "
      "Se il problema persiste, chiudere e riaprire la mini-app dalla voce "
      "Strumenti.")
    h2("Un sito istituzionale non carica")
    p("Alcuni siti del Ministero (es. SIGP) sono di vecchia generazione e "
      "potrebbero apparire piccoli; usare la pinza con due dita per "
      "ingrandire. In caso di lentezza, dopo 5 secondi compare un pulsante "
      "\"Apri in browser\" per usare il browser di sistema.")
    h2("Banner cookie ricorrente")
    p("L'app conserva in memoria persistente i cookie del browser interno: "
      "una volta accettato il banner privacy di un sito, non dovrebbe "
      "ripresentarsi alle aperture successive (la persistenza viene salvata "
      "quando si chiude l'app).")
    h2("Le notifiche PST non arrivano")
    p("Verificare che il toggle in Processo Telematico → Notifiche sia "
      "attivo, che il sistema operativo abbia concesso i permessi notifiche "
      "all'app, e che l'app sia esclusa dall'ottimizzazione batteria del "
      "device (Impostazioni → Batteria → Ottimizzazione → Ordine Avvocati "
      "Napoli → Non ottimizzare).")
    h2("Login Area Riservata: non vedo l'area Consiglieri")
    p("Dopo il login il sito non redirige automaticamente all'area "
      "Consiglieri: torna in app e tocca il bottone blu \"Vai all'Area "
      "Riservata Consiglieri\" (Step 2). Senza questo secondo tocco la "
      "sessione resta attiva nel browser ma non si entra nell'area vera "
      "e propria.")
    h2("Codici Italiani: il testo di un articolo va a capo in modo strano")
    p("I codici sono presentati nello stesso formato della webapp originale "
      "di avvocatotelematico.studiolegalearcella.it: ogni comma in un proprio "
      "paragrafo. Se la formattazione di un articolo specifico risulta "
      "difficile da leggere, si può comunque copiarlo intero col pulsante "
      "\"Copia\" e incollarlo in un editor di testo dove il rendering sarà "
      "ottimale.")

    # ===================================================================
    # FOOTER
    # ===================================================================
    pb()
    C.append(("FOOTER", "— Fine del Manuale —"))
    C.append(("FOOTER_SUB", "App Ordine Avvocati Napoli — Versione 1.1.1 — Documento generato automaticamente"))

    return C


# ============== Render DOCX ==============
# Larghezza standard immagine (cm) — calibrata su screenshot mobile portrait
# (proporzioni ~9:19), per non occupare più della metà di una pagina A4.
IMG_WIDTH_CM = 7.0


def render_docx(content):
    doc = Document()
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Default margini più generosi (Word di default è 2.54 cm: ok)

    def add_heading(text, level):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(*PRIMARY_DARK)
        # Heading deve restare con il paragrafo successivo (no orfana)
        h.paragraph_format.keep_with_next = True
        return h

    def add_para(text, **kw):
        p = doc.add_paragraph()
        # Supporto markdown ** per bold
        parts = text.split("**")
        bold = False
        for part in parts:
            if part:
                r = p.add_run(part)
                r.bold = bold
                if kw.get("italic"):
                    r.italic = True
                if kw.get("color"):
                    r.font.color.rgb = RGBColor(*kw["color"])
                if kw.get("size"):
                    r.font.size = Pt(kw["size"])
            bold = not bold
        return p

    def add_image(path, caption=None, width_cm=IMG_WIDTH_CM):
        if not path or not Path(path).exists():
            return
        # Forza il paragrafo precedente (testo che introduce l'immagine) a
        # restare con quello dell'immagine: evita che il testo finisca a fine
        # pagina e l'immagine inizi nella pagina successiva orfana.
        if doc.paragraphs:
            doc.paragraphs[-1].paragraph_format.keep_with_next = True
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # L'intero blocco immagine + caption non si spezza
        p.paragraph_format.keep_together = True
        p.paragraph_format.keep_with_next = True
        r = p.add_run()
        try:
            r.add_picture(str(path), width=Cm(width_cm))
        except Exception:
            return
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.keep_together = True
            cap.paragraph_format.keep_with_next = True
            cr = cap.add_run(caption)
            cr.italic = True
            cr.font.size = Pt(9)
            cr.font.color.rgb = RGBColor(*MUTED)

    for kind, payload in content:
        if kind == "TITLE":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            r = p.add_run(payload)
            r.font.size = Pt(28)
            r.bold = True
            r.font.color.rgb = RGBColor(*PRIMARY_DARK)
        elif kind == "SUBTITLE":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            r = p.add_run(payload)
            r.font.size = Pt(20)
            r.font.color.rgb = RGBColor(*PRIMARY)
        elif kind == "SUB2":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            r = p.add_run(payload)
            r.italic = True
            r.font.size = Pt(13)
        elif kind == "SUB3":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            r = p.add_run(payload)
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(*MUTED)
            doc.add_paragraph()
        elif kind == "CREDIT":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(payload)
            r.bold = True
            r.font.size = Pt(11)
        elif kind == "CREDIT_SUB":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(payload)
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(*MUTED)
        elif kind == "H1":
            add_heading(payload, level=1)
        elif kind == "H2":
            add_heading(payload, level=2)
        elif kind == "H3":
            add_heading(payload, level=3)
        elif kind == "P":
            add_para(payload)
        elif kind == "BULLET":
            p = doc.add_paragraph(style="List Bullet")
            parts = payload.split("**")
            bold = False
            for part in parts:
                if part:
                    r = p.add_run(part)
                    r.bold = bold
                bold = not bold
        elif kind == "NUM":
            doc.add_paragraph(payload, style="List Number")
        elif kind == "IMG":
            path, caption = payload
            add_image(path, caption)
        elif kind == "NOTE":
            p = doc.add_paragraph()
            p.paragraph_format.keep_together = True
            r = p.add_run("ⓘ  " + payload)
            r.italic = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(*MUTED)
        elif kind == "QUOTE":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.7)
            p.paragraph_format.keep_together = True
            r = p.add_run(payload)
            r.italic = True
        elif kind == "PAGEBREAK":
            doc.add_page_break()
        elif kind == "FOOTER":
            f = doc.add_paragraph()
            f.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = f.add_run(payload)
            r.italic = True
            r.font.color.rgb = RGBColor(*MUTED)
        elif kind == "FOOTER_SUB":
            f = doc.add_paragraph()
            f.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = f.add_run(payload)
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(*MUTED)

    doc.save(str(OUT_DOCX))
    print(f"\n[DOCX] {OUT_DOCX} ({OUT_DOCX.stat().st_size / 1024:.1f} KB)")


# ============== Render Markdown ==============
def render_md(content):
    lines = []
    bullet_open = False
    num_open = False

    def close_lists():
        nonlocal bullet_open, num_open
        if bullet_open or num_open:
            lines.append("")
        bullet_open = False
        num_open = False

    for kind, payload in content:
        if kind in ("H1", "H2", "H3", "P", "IMG", "NOTE", "QUOTE", "PAGEBREAK",
                    "TITLE", "SUBTITLE", "SUB2", "SUB3", "CREDIT", "CREDIT_SUB",
                    "FOOTER", "FOOTER_SUB"):
            close_lists()

        if kind == "TITLE":
            lines.append(f"# {payload}")
            lines.append("")
        elif kind == "SUBTITLE":
            lines.append(f"## {payload}")
            lines.append("")
        elif kind == "SUB2":
            lines.append(f"*{payload}*")
            lines.append("")
        elif kind == "SUB3":
            lines.append(f"*{payload}*")
            lines.append("")
        elif kind == "CREDIT":
            lines.append(f"**{payload}**")
            lines.append("")
        elif kind == "CREDIT_SUB":
            lines.append(f"*{payload}*")
            lines.append("")
        elif kind == "H1":
            lines.append(f"## {payload}")
            lines.append("")
        elif kind == "H2":
            lines.append(f"### {payload}")
            lines.append("")
        elif kind == "H3":
            lines.append(f"#### {payload}")
            lines.append("")
        elif kind == "P":
            lines.append(payload)
            lines.append("")
        elif kind == "BULLET":
            if not bullet_open:
                bullet_open = True
            lines.append(f"- {payload}")
        elif kind == "NUM":
            if not num_open:
                num_open = True
            lines.append(f"- {payload}")
        elif kind == "IMG":
            path, caption = payload
            if path and Path(path).exists():
                rel = "scripts/manual-images/" + Path(path).name
                lines.append(f"![{caption or ''}]({rel})")
                if caption:
                    lines.append("")
                    lines.append(f"*{caption}*")
                lines.append("")
        elif kind == "NOTE":
            lines.append(f"> ℹ️  {payload}")
            lines.append("")
        elif kind == "QUOTE":
            lines.append(f"> {payload}")
            lines.append("")
        elif kind == "PAGEBREAK":
            lines.append("")
            lines.append("---")
            lines.append("")
        elif kind == "FOOTER":
            lines.append(f"*{payload}*")
            lines.append("")
        elif kind == "FOOTER_SUB":
            lines.append(f"*{payload}*")
            lines.append("")

    OUT_MD.write_text("\n".join(lines), encoding="utf-8")
    print(f"[MD]   {OUT_MD} ({OUT_MD.stat().st_size / 1024:.1f} KB)")


if __name__ == "__main__":
    print("Copio screenshot...")
    images = copy_photos()
    print("\nCostruisco contenuto...")
    content = build_content(images)
    print(f"  {len(content)} elementi")
    print("\nGenero DOCX...")
    render_docx(content)
    print("\nGenero MD...")
    render_md(content)
    print("\nFatto.")
